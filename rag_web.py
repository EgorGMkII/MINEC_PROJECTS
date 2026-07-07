from __future__ import annotations

import html
import os
import re
import traceback
from io import BytesIO
from typing import Any

import mammoth
from docx import Document
from fastapi import FastAPI, File, Form, UploadFile
from fastapi.responses import HTMLResponse, JSONResponse, Response, StreamingResponse

from old_prompt import OLD_SYSTEM_PROMPT
from rag_support import (
    DEFAULT_ANALYSIS_TASK,
    DEFAULT_SYSTEM_PROMPT,
    NPA_LABELS,
    list_npa_files,
    parse_upload_bytes,
)
from summary_rag import SimpleRAGPipeline


APP_TITLE = "Проверка изменений в госпрограммы"
PUBLIC_APP_TITLE = "Проект заключения на изменения в госпрограммы"
DEVELOPER_MODE = True
DEFAULT_TOP_K = int(os.getenv("RAG_TOP_K", "20"))
DEFAULT_CHUNK_SIZE = int(os.getenv("RAG_CHUNK_SIZE", "1400"))
DEFAULT_CHUNK_OVERLAP = int(os.getenv("RAG_CHUNK_OVERLAP", "250"))

app = FastAPI(title=APP_TITLE)


def get_app_title() -> str:
    return APP_TITLE if DEVELOPER_MODE else PUBLIC_APP_TITLE


def get_default_system_prompt() -> str:
    return DEFAULT_SYSTEM_PROMPT if DEVELOPER_MODE else OLD_SYSTEM_PROMPT


def render_inline_markup(text: str) -> str:
    escaped = html.escape(text)
    return re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", escaped)


def normalize_answer_text_for_display(text: str) -> str:
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    label_pattern = (
        r"Фрагмент:|Как исправить:|Почему это важно:|"
        r"Что проверить или уточнить:|Ошибка:"
    )
    normalized = re.sub(rf"(?<!\n)\s+({label_pattern})", r"\n\1", normalized)
    normalized = re.sub(r"\s+•\s*", "\n- ", normalized)
    normalized = re.sub(r"(?<=[;:])\s+[–-]\s+", "\n- ", normalized)
    normalized = re.sub(r"(?m)^\s*•\s*", "- ", normalized)
    normalized = re.sub(r"(?m)^\s*[–-]\s+", "- ", normalized)
    normalized = re.sub(r"\n{3,}", "\n\n", normalized)
    return normalized.strip()


def render_answer_html(answer: str) -> str:
    lines = normalize_answer_text_for_display(answer).splitlines()
    html_parts: list[str] = []
    list_items: list[str] = []

    def flush_list() -> None:
        if not list_items:
            return
        html_parts.append("<ul>" + "".join(list_items) + "</ul>")
        list_items.clear()

    for raw_line in lines:
        line = raw_line.strip()
        if not line:
            flush_list()
            continue

        if line.startswith("- "):
            list_items.append(f"<li>{render_inline_markup(line[2:].strip())}</li>")
            continue

        flush_list()
        html_parts.append(f"<p>{render_inline_markup(line)}</p>")

    flush_list()
    return "".join(html_parts)


def render_docx_preview_html(filename: str, content: bytes) -> str:
    fallback_note = ""
    try:
        result = mammoth.convert_to_html(BytesIO(content))
        body_html = result.value or "<p>Документ не содержит отображаемого текста.</p>"
    except Exception:  # noqa: BLE001
        try:
            parsed_text = parse_upload_bytes(filename, content)
        except Exception as exc:  # noqa: BLE001
            parsed_text = f"Не удалось извлечь текст из документа: {exc}"
        body_html = f"<pre>{html.escape(parsed_text)}</pre>"
        result = None
        fallback_note = "<p>HTML-предпросмотр недоступен, показан извлеченный plain text.</p>"
    messages_html = ""
    if result and result.messages:
        items = "".join(f"<li>{html.escape(message.message)}</li>" for message in result.messages)
        messages_html = f"""
        <section class="notes">
          <h2>Замечания конвертации</h2>
          {fallback_note}
          <ul>{items}</ul>
        </section>
        """
    elif fallback_note:
        messages_html = f"""
        <section class="notes">
          <h2>Замечания конвертации</h2>
          {fallback_note}
        </section>
        """
    escaped_name = html.escape(filename)
    return f"""<!doctype html>
<html lang="ru">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>Просмотр: {escaped_name}</title>
  <style>
    body {{
      margin: 0;
      font-family: Calibri, Arial, sans-serif;
      background: #f6f1e8;
      color: #1f1b18;
    }}
    main {{
      max-width: 1274px;
      margin: 0 auto;
      padding: 24px 18px 40px;
    }}
    h1 {{
      margin: 0 0 16px;
      font-size: 1.9rem;
    }}
    .panel, .notes {{
      background: rgba(255, 251, 245, 0.95);
      border: 1px solid rgba(139, 58, 47, 0.16);
      border-radius: 20px;
      padding: 20px;
      box-shadow: 0 14px 40px rgba(73, 50, 25, 0.08);
    }}
    .notes {{
      margin-top: 16px;
    }}
    p, li {{
      font-size: 1rem;
      line-height: 1.6;
    }}
    pre {{
      white-space: pre-wrap;
      word-break: break-word;
      font-family: Calibri, Arial, sans-serif;
      font-size: 1rem;
      line-height: 1.55;
      margin: 0;
    }}
    table {{
      width: 100%;
      border-collapse: collapse;
      margin: 14px 0;
      background: #fff;
    }}
    th, td {{
      border: 1px solid rgba(139, 58, 47, 0.18);
      padding: 8px 10px;
      text-align: left;
      vertical-align: top;
    }}
    h2, h3 {{
      margin-top: 1.3em;
    }}
  </style>
</head>
<body>
  <main>
    <h1>{escaped_name}</h1>
    <section class="panel">{body_html}</section>
    {messages_html}
  </main>
</body>
</html>"""


def render_page(
    answer: str = "",
    error: str = "",
    debug_data: dict[str, Any] | None = None,
    usage_data: dict[str, Any] | None = None,
) -> str:
    app_title = get_app_title()
    npa_items = "".join(
        f"<li>{html.escape(NPA_LABELS.get(name, name))}</li>"
        for name in list_npa_files()
    )
    hidden_corpus_items = npa_items + "<li>Замечания УФАС по НСО</li>"

    retrieved_items = ""
    if debug_data:
        for document in debug_data.get("documents", []):
            metadata = document.metadata
            retrieved_items += (
                "<li>"
                f"{html.escape(str(metadata.get('group_label', metadata.get('group', 'UNKNOWN'))))} / "
                f"{html.escape(str(metadata.get('title', metadata.get('source', 'unknown'))))} / "
                f"чанк {html.escape(str(metadata.get('chunk_order', '?')))}"
                "</li>"
            )

    answer_html = ""
    if answer:
        answer_html = render_answer_html(answer)

    error_html = f"<div class='error'>{html.escape(error)}</div>" if error else ""

    debug_html = ""
    if debug_data:
        retrieval_query = html.escape(debug_data.get("retrieval_query", ""))
        debug_html = f"""
        <section class="panel">
          <h2>Что попало в retrieval</h2>
          <details>
            <summary>Показать retrieval query</summary>
            <pre>{retrieval_query}</pre>
          </details>
          <h3>Найденные чанки</h3>
          <ul>{retrieved_items or "<li>Ничего не найдено</li>"}</ul>
        </section>
        """

    download_html = ""
    if answer:
        download_html = f"""
        <form action="/download-docx" method="post" style="margin: 0 0 14px;">
          <input type="hidden" name="answer_text" value="{html.escape(answer, quote=True)}">
          <button type="submit">Скачать ответ в DOCX</button>
        </form>
        """

    usage_html = ""
    if usage_data:
        total_tokens = usage_data.get("total_tokens")
        input_tokens = usage_data.get("input_tokens")
        output_tokens = usage_data.get("output_tokens")
        usage_parts: list[str] = []
        if total_tokens is not None:
            usage_parts.append(f"Всего токенов: {total_tokens}")
        if input_tokens is not None:
            usage_parts.append(f"Вход: {input_tokens}")
        if output_tokens is not None:
            usage_parts.append(f"Выход: {output_tokens}")
        if usage_parts:
            usage_html = f"<p class='usage'>{html.escape(' | '.join(usage_parts))}</p>"

    return f"""<!doctype html>
<html lang="ru">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>{html.escape(app_title)}</title>
  <style>
    :root {{
      --panel: rgba(255, 251, 245, 0.92);
      --ink: #1f1b18;
      --muted: #6a5f57;
      --line: #d9cab7;
      --accent: #8b3a2f;
      --error: #b42318;
    }}
    * {{ box-sizing: border-box; }}
    body {{
      margin: 0;
      color: var(--ink);
      background:
        radial-gradient(circle at top left, rgba(209, 161, 93, 0.28), transparent 32%),
        radial-gradient(circle at bottom right, rgba(139, 58, 47, 0.15), transparent 28%),
        linear-gradient(135deg, #f7f2ea 0%, #efe4d4 100%);
      font-family: Calibri, Arial, sans-serif;
    }}
    .shell {{
      max-width: 1534px;
      margin: 0 auto;
      padding: 28px 18px 36px;
    }}
    .hero {{
      padding: 24px 24px 18px;
      border: 1px solid rgba(139, 58, 47, 0.18);
      background: linear-gradient(180deg, rgba(255,255,255,0.7), rgba(255,255,255,0.45));
      backdrop-filter: blur(8px);
      border-radius: 24px;
      box-shadow: 0 18px 54px rgba(73, 50, 25, 0.10);
    }}
    h1, h2, h3 {{
      margin: 0 0 12px;
      font-weight: 700;
      letter-spacing: 0.01em;
    }}
    h1 {{ font-size: clamp(2rem, 4vw, 3.1rem); }}
    h2 {{ font-size: 1.35rem; }}
    h3 {{ font-size: 1.05rem; }}
    p, li, label, input, button, summary, a {{
      font-size: 1rem;
      line-height: 1.55;
    }}
    .lead {{
      max-width: 78ch;
      color: var(--muted);
      margin-bottom: 0;
    }}
    .grid {{
      display: grid;
      grid-template-columns: minmax(340px, 460px) minmax(0, 1fr);
      gap: 18px;
      margin-top: 18px;
    }}
    .panel {{
      background: var(--panel);
      border: 1px solid rgba(139, 58, 47, 0.16);
      border-radius: 22px;
      padding: 20px;
      box-shadow: 0 14px 40px rgba(73, 50, 25, 0.08);
    }}
    .field {{
      margin-bottom: 16px;
    }}
    .field-head {{
      display: flex;
      justify-content: space-between;
      align-items: center;
      gap: 10px;
      margin-bottom: 7px;
    }}
    .field-tools {{
      display: flex;
      align-items: center;
      gap: 8px;
      flex-wrap: wrap;
    }}
    .field label {{
      display: block;
      font-weight: 700;
    }}
    .field input {{
      width: 100%;
      padding: 12px 14px;
      border: 1px solid var(--line);
      border-radius: 14px;
      background: rgba(255,255,255,0.86);
      color: var(--ink);
    }}
    .preview,
    .preview-list {{
      margin-top: 8px;
      color: var(--muted);
      font-size: 0.95rem;
    }}
    .preview button,
    .preview-list button {{
      color: var(--accent);
      background: none;
      border: none;
      padding: 0;
      width: auto;
      box-shadow: none;
      border-bottom: 1px dotted rgba(139, 58, 47, 0.4);
      border-radius: 0;
      cursor: pointer;
    }}
    .preview button:hover,
    .preview-list button:hover {{
      border-bottom-style: solid;
      transform: none;
    }}
    .slot-clear {{
      width: auto;
      padding: 7px 12px;
      border-radius: 999px;
      background: rgba(255,255,255,0.82);
      color: var(--accent);
      border: 1px solid rgba(139, 58, 47, 0.18);
      box-shadow: none;
      font-weight: 700;
      font-size: 0.92rem;
    }}
    .slot-clear:hover {{
      transform: none;
      background: rgba(255,255,255,0.95);
    }}
    .preview-list ul {{
      margin-top: 8px;
      padding-left: 0;
      list-style: none;
    }}
    .preview-list li {{
      display: flex;
      align-items: center;
      gap: 8px;
      margin-bottom: 6px;
    }}
    .preview-remove {{
      width: auto;
      min-width: 24px;
      padding: 0;
      border: none;
      background: none;
      color: var(--accent);
      box-shadow: none;
      font-weight: 700;
      font-size: 1rem;
      line-height: 1;
      cursor: pointer;
    }}
    .preview-remove:hover {{
      transform: none;
      color: #7f2f23;
    }}
    .preview-open {{
      text-align: left;
    }}
    .hint {{
      color: var(--muted);
      margin-top: -2px;
      margin-bottom: 12px;
    }}
    .actions {{
      display: flex;
      gap: 12px;
      margin-top: 6px;
    }}
    button {{
      width: 100%;
      border: 0;
      border-radius: 999px;
      padding: 14px 18px;
      color: #fff8f0;
      background: linear-gradient(135deg, var(--accent), #b45c33);
      font-weight: 700;
      cursor: pointer;
      box-shadow: 0 10px 22px rgba(139, 58, 47, 0.18);
    }}
    button:hover {{
      transform: translateY(-1px);
    }}
    .ghost {{
      width: 100%;
      display: inline-flex;
      align-items: center;
      justify-content: center;
      border-radius: 999px;
      padding: 14px 18px;
      color: var(--accent);
      background: rgba(255,255,255,0.76);
      border: 1px solid rgba(139, 58, 47, 0.18);
      font-weight: 700;
      text-decoration: none;
    }}
    ul {{
      padding-left: 20px;
      margin: 0;
    }}
    pre {{
      white-space: pre-wrap;
      word-break: break-word;
      background: rgba(255,255,255,0.78);
      border: 1px solid rgba(139, 58, 47, 0.12);
      border-radius: 14px;
      padding: 14px;
      max-height: 320px;
      overflow: auto;
    }}
    .answer p {{
      margin: 0 0 12px;
    }}
    .answer ul {{
      margin: 0 0 12px 22px;
      padding-left: 18px;
    }}
    .answer li {{
      margin-bottom: 6px;
    }}
    .usage {{
      margin: 0 0 12px;
      color: var(--muted);
      font-size: 0.95rem;
    }}
    .error {{
      margin: 0 0 16px;
      padding: 12px 14px;
      border-radius: 14px;
      color: var(--error);
      background: rgba(180, 35, 24, 0.08);
      border: 1px solid rgba(180, 35, 24, 0.22);
    }}
    .meta {{
      display: grid;
      gap: 18px;
    }}
    .loading-overlay {{
      position: fixed;
      inset: 0;
      display: none;
      align-items: center;
      justify-content: center;
      background: rgba(247, 242, 234, 0.82);
      backdrop-filter: blur(4px);
      z-index: 9999;
    }}
    .loading-overlay.active {{
      display: flex;
    }}
    .loading-card {{
      min-width: 280px;
      max-width: 420px;
      padding: 24px 22px;
      border-radius: 22px;
      background: rgba(255, 251, 245, 0.98);
      border: 1px solid rgba(139, 58, 47, 0.16);
      box-shadow: 0 18px 54px rgba(73, 50, 25, 0.12);
      text-align: center;
    }}
    .spinner {{
      width: 52px;
      height: 52px;
      margin: 0 auto 16px;
      border-radius: 50%;
      border: 5px solid rgba(139, 58, 47, 0.16);
      border-top-color: var(--accent);
      animation: spin 0.95s linear infinite;
    }}
    .loading-title {{
      margin: 0 0 8px;
      font-size: 1.08rem;
      font-weight: 700;
    }}
    .loading-text {{
      margin: 0;
      color: var(--muted);
    }}
    @keyframes spin {{
      to {{
        transform: rotate(360deg);
      }}
    }}
    @media (max-width: 940px) {{
      .grid {{
        grid-template-columns: 1fr;
      }}
      .actions {{
        flex-direction: column;
      }}
      .shell {{
        padding: 18px 14px 26px;
      }}
      .hero,
      .panel {{
        border-radius: 18px;
      }}
    }}
  </style>
</head>
<body>
  <div id="loading-overlay" class="loading-overlay" aria-live="polite" aria-hidden="true">
    <div class="loading-card">
      <div class="spinner"></div>
      <p class="loading-title">Файлы обрабатываются</p>
      <p class="loading-text">Суммаризация и поиск релевантного контекста могут занять некоторое время.</p>
    </div>
  </div>
  <main class="shell">
    <section class="hero">
      <h1>{html.escape(app_title)}</h1>
        <p class="lead">
        Загрузите пакет проекта изменений, включая основной проект и приложения, региональную программу, федеральную программу и дополнительные документы.
        Нормативная база лежит внутри сервиса и автоматически участвует в retrieval.
      </p>
    </section>

    <section class="grid">
      <form id="analyze-form" class="panel" action="/analyze" method="post" enctype="multipart/form-data">
        <h2>Запуск анализа</h2>
        {error_html}
        <div class="field">
          <div class="field-head">
            <label for="project_changes_file">Проект изменений (проект + приложения)</label>
            <div class="field-tools">
              <button type="button" class="slot-clear" data-clear-slot="project_changes_file" data-preview-target="project_changes_file_preview">Очистить</button>
            </div>
          </div>
          <input id="project_changes_file" name="project_changes_file" type="file" accept=".docx,.pdf,.txt,.xlsx" multiple required>
          <div id="project_changes_file_preview" class="preview-list"></div>
        </div>
        <div class="field">
          <div class="field-head">
            <label for="regional_program_file">Региональная программа</label>
            <div class="field-tools">
              <button type="button" class="slot-clear" data-clear-slot="regional_program_file" data-preview-target="regional_program_file_preview">Очистить</button>
            </div>
          </div>
          <input id="regional_program_file" name="regional_program_file" type="file" accept=".docx,.pdf,.txt,.xlsx" required>
          <div id="regional_program_file_preview" class="preview"></div>
        </div>
        <div class="field">
          <div class="field-head">
            <label for="federal_program_file">Федеральная программа</label>
            <div class="field-tools">
              <button type="button" class="slot-clear" data-clear-slot="federal_program_file" data-preview-target="federal_program_file_preview">Очистить</button>
            </div>
          </div>
          <input id="federal_program_file" name="federal_program_file" type="file" accept=".docx,.pdf,.txt,.xlsx">
          <div id="federal_program_file_preview" class="preview"></div>
        </div>
        <div class="field">
          <div class="field-head">
            <label for="extra_documents">Дополнительные документы</label>
            <div class="field-tools">
              <button type="button" class="slot-clear" data-clear-slot="extra_documents" data-preview-target="extra_documents_preview">Очистить</button>
            </div>
          </div>
          <input id="extra_documents" name="extra_documents" type="file" accept=".docx,.pdf,.txt,.xlsx" multiple>
          <div id="extra_documents_preview" class="preview-list"></div>
        </div>
        <div class="actions">
          <button type="submit">Сформировать заключение</button>
          <a class="ghost" href="/" id="clear-all">Очистить</a>
        </div>
      </form>

      <div class="meta">
        <section class="panel">
          <h2>Скрытый корпус НПА</h2>
          <p class="hint">Эти документы лежат внутри сервиса и автоматически участвуют в BM25 retrieval.</p>
          <ul>{hidden_corpus_items}</ul>
        </section>

        <section id="answer-panel" class="panel answer">
          <h2>Ответ модели</h2>
          {download_html}
          {usage_html}
          {answer_html or "<p class='hint'>Пока пусто. После запуска здесь появится заключение.</p>"}
        </section>
        {debug_html}
      </div>
    </section>
  </main>
  <script>
    const singlePreviewIds = [
      ["regional_program_file", "regional_program_file_preview"],
      ["federal_program_file", "federal_program_file_preview"],
    ];
    const multiFileState = {{
      project_changes_file: [],
      extra_documents: [],
    }};

    function renderSinglePreview(inputId, previewId) {{
      const input = document.getElementById(inputId);
      const preview = document.getElementById(previewId);
      if (!input || !preview) return;

      input.addEventListener("change", () => {{
        preview.innerHTML = "";
        const file = input.files && input.files[0];
        if (!file) return;

        const url = URL.createObjectURL(file);
        preview.innerHTML = `<button type="button" data-preview-single="${{inputId}}">Открыть: ${{file.name}}</button>`;
      }});
    }}

    function syncMultiInputFiles(inputId) {{
      const input = document.getElementById(inputId);
      if (!input) return;
      const dt = new DataTransfer();
      for (const file of multiFileState[inputId] || []) {{
        dt.items.add(file);
      }}
      input.files = dt.files;
    }}

    function mergeFiles(inputId, newFiles) {{
      const existing = multiFileState[inputId] || [];
      const merged = [...existing];
      for (const file of newFiles) {{
        const duplicate = merged.some(
          (current) =>
            current.name === file.name &&
            current.size === file.size &&
            current.lastModified === file.lastModified,
        );
        if (!duplicate) {{
          merged.push(file);
        }}
      }}
      multiFileState[inputId] = merged;
      syncMultiInputFiles(inputId);
    }}

    function clearFiles(inputId) {{
      const input = document.getElementById(inputId);
      multiFileState[inputId] = [];
      if (input && input.multiple) {{
        syncMultiInputFiles(inputId);
      }} else if (input) {{
        input.value = "";
      }}
    }}

    function updateMultiPreview(inputId, previewId, dataKey) {{
      const preview = document.getElementById(previewId);
      if (!preview) return;
      preview.innerHTML = "";
      const files = multiFileState[inputId] || [];
      if (!files.length) return;

      const list = document.createElement("ul");
      files.forEach((file, index) => {{
        const item = document.createElement("li");
        const removeButton = document.createElement("button");
        removeButton.type = "button";
        removeButton.className = "preview-remove";
        removeButton.dataset.removeFile = inputId;
        removeButton.dataset.removeIndex = String(index);
        removeButton.setAttribute("aria-label", `Удалить ${{file.name}}`);
        removeButton.textContent = "×";
        const button = document.createElement("button");
        button.type = "button";
        button.className = "preview-open";
        button.dataset[dataKey] = String(index);
        button.textContent = `Открыть: ${{file.name}}`;
        item.appendChild(removeButton);
        item.appendChild(button);
        list.appendChild(item);
      }});
      preview.appendChild(list);
    }}

    function removeMultiFile(inputId, index, previewId, dataKey) {{
      const files = multiFileState[inputId] || [];
      if (index < 0 || index >= files.length) return;
      multiFileState[inputId] = files.filter((_, currentIndex) => currentIndex !== index);
      syncMultiInputFiles(inputId);
      updateMultiPreview(inputId, previewId, dataKey);
    }}

    function renderMultiPreview(inputId, previewId, dataKey) {{
      const input = document.getElementById(inputId);
      const preview = document.getElementById(previewId);
      if (!input || !preview) return;

      input.addEventListener("change", () => {{
        const files = Array.from(input.files || []);
        if (!files.length) return;
        mergeFiles(inputId, files);
        updateMultiPreview(inputId, previewId, dataKey);
      }});

      input.addEventListener("dragover", (event) => {{
        event.preventDefault();
      }});

      input.addEventListener("drop", (event) => {{
        event.preventDefault();
        const files = Array.from(event.dataTransfer?.files || []);
        if (!files.length) return;
        mergeFiles(inputId, files);
        updateMultiPreview(inputId, previewId, dataKey);
      }});
    }}

    async function openPreview(file) {{
      if (!file) return;
      const formData = new FormData();
      formData.append("file", file);
      const response = await fetch("/preview", {{
        method: "POST",
        body: formData,
      }});
      if (!response.ok) {{
        const text = await response.text();
        alert(text || "Не удалось открыть предпросмотр.");
        return;
      }}
      const htmlText = await response.text();
      const previewWindow = window.open("", "_blank");
      if (!previewWindow) {{
        alert("Браузер заблокировал новое окно предпросмотра.");
        return;
      }}
      previewWindow.document.open();
      previewWindow.document.write(htmlText);
      previewWindow.document.close();
    }}

    for (const [inputId, previewId] of singlePreviewIds) {{
      renderSinglePreview(inputId, previewId);
    }}
    renderMultiPreview("project_changes_file", "project_changes_file_preview", "previewProjectChangesFile");
    renderMultiPreview("extra_documents", "extra_documents_preview", "previewExtra");

    document.addEventListener("click", async (event) => {{
      const clearButton = event.target.closest("[data-clear-slot]");
      if (clearButton) {{
        const inputId = clearButton.getAttribute("data-clear-slot");
        const previewId = clearButton.getAttribute("data-preview-target");
        clearFiles(inputId);
        const preview = previewId ? document.getElementById(previewId) : null;
        if (preview) preview.innerHTML = "";
        return;
      }}

      const singleButton = event.target.closest("[data-preview-single]");
      if (singleButton) {{
        const inputId = singleButton.getAttribute("data-preview-single");
        const input = document.getElementById(inputId);
        const file = input && input.files ? input.files[0] : null;
        await openPreview(file);
        return;
      }}

      const removeButton = event.target.closest("[data-remove-file]");
      if (removeButton) {{
        const inputId = removeButton.getAttribute("data-remove-file");
        const index = Number(removeButton.getAttribute("data-remove-index"));
        if (inputId === "project_changes_file") {{
          removeMultiFile("project_changes_file", index, "project_changes_file_preview", "previewProjectChangesFile");
        }} else if (inputId === "extra_documents") {{
          removeMultiFile("extra_documents", index, "extra_documents_preview", "previewExtra");
        }}
        return;
      }}

      const extraButton = event.target.closest("[data-preview-extra]");
      if (extraButton) {{
        const index = Number(extraButton.getAttribute("data-preview-extra"));
        const input = document.getElementById("extra_documents");
        const file = input && input.files ? input.files[index] : null;
        await openPreview(file);
        return;
      }}

      const projectButton = event.target.closest("[data-preview-project-changes-file]");
      if (projectButton) {{
        const index = Number(projectButton.getAttribute("data-preview-project-changes-file"));
        const input = document.getElementById("project_changes_file");
        const file = input && input.files ? input.files[index] : null;
        await openPreview(file);
      }}
    }});

    const clearAll = document.getElementById("clear-all");
    const analyzeForm = document.getElementById("analyze-form");
    const loadingOverlay = document.getElementById("loading-overlay");
    if (clearAll && analyzeForm) {{
      clearAll.addEventListener("click", (event) => {{
        event.preventDefault();
        analyzeForm.reset();
        for (const [, previewId] of singlePreviewIds) {{
          const preview = document.getElementById(previewId);
          if (preview) preview.innerHTML = "";
        }}
        const extraPreview = document.getElementById("extra_documents_preview");
        if (extraPreview) extraPreview.innerHTML = "";
        clearFiles("project_changes_file");
        clearFiles("extra_documents");
        window.location.href = "/";
      }});

      analyzeForm.addEventListener("submit", () => {{
        if (loadingOverlay) {{
          loadingOverlay.classList.add("active");
          loadingOverlay.setAttribute("aria-hidden", "false");
        }}
      }});
    }}
  </script>
</body>
</html>"""


async def build_pipeline_from_uploads(
    regional_program_file: UploadFile,
    project_changes_files: list[UploadFile],
    federal_program_file: UploadFile | None = None,
    extra_documents: list[UploadFile] | None = None,
    system_prompt: str | None = None,
) -> SimpleRAGPipeline:
    effective_system_prompt = system_prompt if DEVELOPER_MODE else None
    regional_program_bytes = await regional_program_file.read()
    project_files: list[tuple[str, bytes]] = []

    federal_files: list[tuple[str, bytes]] = []
    extra_files: list[tuple[str, bytes]] = []

    for project_changes_file in project_changes_files:
        if project_changes_file.filename:
            project_files.append(
                (
                    project_changes_file.filename or "project_changes.docx",
                    await project_changes_file.read(),
                )
            )

    if federal_program_file and federal_program_file.filename:
        federal_files.append(
            (
                federal_program_file.filename or "federal_program.docx",
                await federal_program_file.read(),
            )
        )

    for extra_document in extra_documents or []:
        if extra_document.filename:
            extra_files.append(
                (
                    extra_document.filename or "extra_document.docx",
                    await extra_document.read(),
                )
            )

    return SimpleRAGPipeline.from_uploads(
        project_files=project_files,
        base_program_files=[(regional_program_file.filename or "regional_program.docx", regional_program_bytes)],
        npa_files=list_npa_files(),
        federal_files=federal_files,
        extra_files=extra_files,
        top_k=DEFAULT_TOP_K,
        chunk_size=DEFAULT_CHUNK_SIZE,
        chunk_overlap=DEFAULT_CHUNK_OVERLAP,
        base_system_prompt=effective_system_prompt or get_default_system_prompt(),
    )


def serialize_debug(debug_data: dict[str, Any]) -> dict[str, Any]:
    return {
        "retrieval_query": debug_data.get("retrieval_query", ""),
        "context": debug_data.get("context", ""),
        "documents": [
            {
                "page_content": document.page_content,
                "metadata": dict(document.metadata),
            }
            for document in debug_data.get("documents", [])
        ],
    }


@app.get("/", response_class=HTMLResponse)
async def index() -> str:
    return render_page()


@app.get("/favicon.ico")
async def favicon() -> Response:
    return Response(status_code=204)


@app.get("/health")
async def health() -> dict[str, Any]:
    return {
        "status": "ok",
        "npa_count": len(list_npa_files()),
        "model": os.getenv("OPENAI_MODEL", "gpt-5.3-chat-latest"),
    }


@app.post("/api/analyze")
async def api_analyze(
    project_changes_file: list[UploadFile] = File(...),
    regional_program_file: UploadFile = File(...),
    federal_program_file: UploadFile | None = File(None),
    extra_documents: list[UploadFile] | None = File(None),
    analysis_task: str = Form(DEFAULT_ANALYSIS_TASK),
    system_prompt: str | None = Form(None),
) -> JSONResponse:
    effective_system_prompt = system_prompt if DEVELOPER_MODE else None
    pipeline = await build_pipeline_from_uploads(
        regional_program_file=regional_program_file,
        project_changes_files=project_changes_file,
        federal_program_file=federal_program_file,
        extra_documents=extra_documents,
        system_prompt=effective_system_prompt,
    )
    debug_data = pipeline.debug(analysis_task)
    answer = debug_data["answer"]
    return JSONResponse(
        {
            "answer": answer,
            "analysis_task": analysis_task,
            "project_text": pipeline.project_text,
            "system_prompt": effective_system_prompt,
            "usage": debug_data["usage"],
            "debug": serialize_debug(debug_data),
        }
    )


@app.post("/analyze", response_class=HTMLResponse)
async def analyze(
    project_changes_file: list[UploadFile] = File(...),
    regional_program_file: UploadFile = File(...),
    federal_program_file: UploadFile | None = File(None),
    extra_documents: list[UploadFile] | None = File(None),
    system_prompt: str | None = Form(None),
) -> str:
    effective_system_prompt = system_prompt if DEVELOPER_MODE else None
    try:
        pipeline = await build_pipeline_from_uploads(
            regional_program_file=regional_program_file,
            project_changes_files=project_changes_file,
            federal_program_file=federal_program_file,
            extra_documents=extra_documents,
            system_prompt=effective_system_prompt,
        )
        debug_data = pipeline.debug(DEFAULT_ANALYSIS_TASK)
        answer = debug_data["answer"]
        return render_page(
            answer=answer,
            debug_data=debug_data,
            usage_data=debug_data["usage"],
        )
    except Exception as exc:  # noqa: BLE001
        traceback.print_exc()
        return render_page(
            error=f"Не удалось выполнить анализ: {exc}",
        )


@app.post("/download-docx")
async def download_docx(answer_text: str = Form(...)) -> StreamingResponse:
    document = Document()
    document.add_heading("Заключение по проекту изменений", level=1)
    for block in answer_text.split("\n\n"):
        cleaned = block.strip()
        if cleaned:
            document.add_paragraph(cleaned)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": 'attachment; filename="answer.docx"'},
    )


@app.post("/preview", response_class=HTMLResponse)
async def preview_file(file: UploadFile = File(...)) -> str:
    content = await file.read()
    escaped_name = html.escape(file.filename or "document")
    filename = file.filename or "document"
    if filename.lower().endswith(".docx"):
        return render_docx_preview_html(filename, content)

    parsed_text = parse_upload_bytes(filename, content)
    escaped_text = html.escape(parsed_text)
    return f"""<!doctype html>
<html lang="ru">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>Просмотр: {escaped_name}</title>
  <style>
    body {{
      margin: 0;
      font-family: Calibri, Arial, sans-serif;
      background: #f6f1e8;
      color: #1f1b18;
    }}
    main {{
      max-width: 1274px;
      margin: 0 auto;
      padding: 24px 18px 40px;
    }}
    h1 {{
      margin: 0 0 16px;
      font-size: 1.9rem;
    }}
    .panel {{
      background: rgba(255, 251, 245, 0.95);
      border: 1px solid rgba(139, 58, 47, 0.16);
      border-radius: 20px;
      padding: 20px;
      box-shadow: 0 14px 40px rgba(73, 50, 25, 0.08);
    }}
    pre {{
      white-space: pre-wrap;
      word-break: break-word;
      font-family: Calibri, Arial, sans-serif;
      font-size: 1rem;
      line-height: 1.55;
      margin: 0;
    }}
  </style>
</head>
<body>
  <main>
    <h1>{escaped_name}</h1>
    <section class="panel">
      <pre>{escaped_text}</pre>
    </section>
  </main>
</body>
</html>"""
